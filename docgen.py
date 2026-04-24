"""docgen.py - 문서 자동 생성 모듈 (템플릿 기반)"""
import io, re, calendar, zipfile, os
from copy import copy
from datetime import date, timedelta

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), 'doc_templates')

MONTHLY_SALARY    = 8_850_000
PM_RATE           = 0.01
PROD_RATE         = 0.15
STUDIO_UNIT_PRICE = 45_000
TARGET_PROFIT     = 0.30
PRICE_NEW       = 500_000
PRICE_PORTING   = 50_000
PRICE_EDIT_PORT = 160_000
PRICE_TRAVEL_HR = 100_000

# 유형별 1차시당 표준 작업시간 (촬영·편집 담당, 시간 단위)
WORK_HOURS_PER_SESSION = {
    'default':      2.5,   # 크로마키·태블릿형·전자칠판형·신규개발
    'porting':      0.5,   # 포팅(무편집)
    'edit_porting': 1.0,   # 포팅(편집)
    'travel':       3.5,   # 출장 (개발 2.5 + 출장비 1.0시간/차시)
}
WORK_HOURS_PER_DAY = 8.0

PM_NAME        = "이상현"
PROD_NAME      = "염왕도"
DEPT_CODE      = "SS"
DEPT_FULL      = "에듀테크서비스본부"
DELIVERY_PLACE = "한국공인회계사회"

def get_last_business_day(year, month):
    d = date(year, month, calendar.monthrange(year, month)[1])
    while d.weekday() >= 5:
        d -= timedelta(days=1)
    return d

def get_next_month_5th_weekday(year, month):
    ny, nm = (year + 1, 1) if month == 12 else (year, month + 1)
    d = date(ny, nm, 5)
    while d.weekday() >= 5:
        d += timedelta(days=1)
    return d

def get_month_number(month_str):
    m = re.search(r'(\d{1,2})월', str(month_str or ""))
    return int(m.group(1)) if m else 1

def fmt_kr(d):    return f"{d.year}. {d.month:02d}. {d.day:02d}"
def fmt_kr2(d):   return f"{d.year}년 {d.month:02d}월 {d.day:02d}일"
def fmt_short(d): return f"{str(d.year)[2:]}.{d.month:02d}.{d.day:02d}"

def get_weekday_count(start, end):
    """start~end 사이 평일(월~금) 수"""
    count, d = 0, start
    while d <= end:
        if d.weekday() < 5:
            count += 1
        d += timedelta(days=1)
    return count

def get_sijengil(ps):
    d = ps - timedelta(days=2)
    while d.weekday() >= 5:
        d -= timedelta(days=1)
    return d

def classify_fmt(fmt):
    if not fmt: return True, False, False
    if "포팅" in fmt:
        if "편집" in fmt and "무편집" not in fmt: return False, True, True
        return False, True, False
    return True, False, False

def get_unit_price_for(cr, price_tbl):
    if cr.custom_price: return cr.custom_price
    fmt = cr.shooting_format or ""
    if "포팅" in fmt:
        if "편집" in fmt and "무편집" not in fmt:
            return price_tbl.get("편집포팅", PRICE_EDIT_PORT) or PRICE_EDIT_PORT
        return price_tbl.get("포팅", PRICE_PORTING) or PRICE_PORTING
    if "출장" in fmt:
        return price_tbl.get("FullVod (출장)", PRICE_NEW) or PRICE_NEW
    for k, v in price_tbl.items():
        if k and k in fmt: return v or 0
    return PRICE_NEW

def get_travel_for(cr, travel_rate=PRICE_TRAVEL_HR):
    """출장비 계산: 1일 4시간 초과 시 4시간으로 캡 적용"""
    if not cr.shooting_format or "출장" not in cr.shooting_format: return 0
    if cr.travel_expense is not None: return cr.travel_expense
    if cr.travel_hours:
        hours = cr.travel_hours
        days  = getattr(cr, 'travel_days', None)
        capped = min(hours, days * 4) if days and days > 0 else hours
        return capped * travel_rate
    return travel_rate

def calc_period(courses, year, month_num):
    starts = [c.shooting_date for c in courses if c.shooting_date]
    ends   = [c.open_date     for c in courses if c.open_date]
    s = min(starts) if starts else date(year, month_num, 1)
    e = max(ends)   if ends   else get_next_month_5th_weekday(year, month_num)
    return s, e

def calc_revenue(courses, price_tbl):
    tr = price_tbl.get("1 ~ 4시간", PRICE_TRAVEL_HR)
    return sum(
        get_unit_price_for(c, price_tbl) * (c.session_count or c.chapter_count or 0)
        + get_travel_for(c, tr) for c in courses)

def calc_labor_amounts(ps, pe, pm_rate, prod_rate):
    d = (pe - ps).days
    return (round(d/30*pm_rate*MONTHLY_SALARY+1),
            round(d/30*prod_rate*MONTHLY_SALARY+1))

def calc_prod_rate_from_standards(courses, ps, pe):
    """콘텐츠 유형별 표준시간 기반 PROD 투입비율 계산"""
    total_work = 0.0
    for c in courses:
        fmt = c.shooting_format or ""
        sessions = c.session_count or 0
        if "포팅" in fmt:
            hours = WORK_HOURS_PER_SESSION['edit_porting'] if ("편집" in fmt and "무편집" not in fmt) else WORK_HOURS_PER_SESSION['porting']
        elif "출장" in fmt:
            hours = WORK_HOURS_PER_SESSION['travel']
        else:
            hours = WORK_HOURS_PER_SESSION['default']
        total_work += sessions * hours

    available = get_weekday_count(ps, pe) * WORK_HOURS_PER_DAY
    if available <= 0:
        return PROD_RATE
    return total_work / available


def adjust_rates(revenue, studio_a, ps, pe, courses=None):
    """PM=1% 고정, PROD=표준시간 기반 계산, 손익률 30% 이상 보장. 비율 정수% 반올림."""
    pm_fixed = 0.01  # PM 1% 고정

    # PROD: 표준시간 기반 계산 → 정수% 반올림
    raw_prod  = calc_prod_rate_from_standards(courses, ps, pe) if courses else PROD_RATE
    prod_rate = round(raw_prod * 100) / 100   # e.g. 0.3125 → 0.31 (31%)

    # 손익률 30% 확인
    pm_a, prod_a = calc_labor_amounts(ps, pe, pm_fixed, prod_rate)
    if revenue > 0 and (revenue - pm_a - prod_a - studio_a) / revenue >= TARGET_PROFIT:
        return pm_fixed, prod_rate  # 30% 이상 → 그대로 사용

    # PROD 축소하여 30% 보장
    period_d  = (pe - ps).days
    max_labor = (1 - TARGET_PROFIT) * revenue - studio_a - 2
    pm_amt    = round(period_d / 30 * pm_fixed * MONTHLY_SALARY + 1)
    max_prod  = max_labor - pm_amt
    if max_prod <= 0 or period_d <= 0:
        return pm_fixed, 0.0
    new_prod  = round(max_prod / (period_d / 30 * MONTHLY_SALARY) * 100) / 100
    return pm_fixed, max(new_prod, 0.0)


def build_project_name(courses):
    ns  = sum(c.session_count or 0 for c in courses if classify_fmt(c.shooting_format or "")[0])
    pc  = sum(c.chapter_count or 0 for c in courses
              if classify_fmt(c.shooting_format or "")[1] and not classify_fmt(c.shooting_format or "")[2])
    ec  = sum(c.chapter_count or 0 for c in courses if classify_fmt(c.shooting_format or "")[2])
    pts = []
    if ns: pts.append(f"신규{ns}차시")
    if pc: pts.append(f"포팅{pc}챕터")
    if ec: pts.append(f"편집포팅{ec}챕터")
    return "한국공인회계사 콘텐츠 개발(" + ("·".join(pts) if pts else "신규") + ")"

def _replace_para(para, old, new):
    full = para.text
    if old not in full: return False
    replaced = full.replace(old, new)
    if para.runs:
        para.runs[0].text = replaced
        for r in para.runs[1:]: r.text = ""
    else:
        para.add_run(replaced)
    return True

def _replace_doc(doc, old, new):
    for p in doc.paragraphs: _replace_para(p, old, new)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: _replace_para(p, old, new)

def _set_cell_text(cell, txt):
    for para in cell.paragraphs:
        if para.runs:
            para.runs[0].text = txt
            for r in para.runs[1:]: r.text = ""
            return
    cell.paragraphs[0].add_run(txt)

def _clear_cell_and_set(cell, txt):
    """셀의 기존 단락을 모두 제거하고 새 텍스트 단락으로 교체"""
    tc = cell._tc
    p_elems = list(tc.findall(qn("w:p")))
    # 두 번째 이후 단락 XML 요소 제거 (템플릿 잔여 단락 삭제)
    for pe in p_elems[1:]:
        tc.remove(pe)
    # 첫 단락 런 초기화 후 새 텍스트 설정
    first = cell.paragraphs[0]
    for r in first.runs: r.text = ""
    if first.runs:
        first.runs[0].text = txt
    else:
        first.add_run(txt)


# 1. 손익분석서
def gen_pnl_excel(courses, dept, month_str, year, price_tbl,
                  studio_hours, include_studio, pm_rate, prod_rate):
    mn = get_month_number(month_str)
    ps, pe = calc_period(courses, year, mn)
    wb = openpyxl.load_workbook(os.path.join(TEMPLATE_DIR, 'tpl_pnl.xlsx'))
    ws = wb.active
    ws['F3'] = build_project_name(courses)
    ws['F4'] = get_last_business_day(year, mn)
    ws['F4'].number_format = 'YYYY"년" MM"월" DD"일"'
    ws['I8'] = calc_revenue(courses, price_tbl)
    ws['D10'] = ps; ws['D10'].number_format = 'YYYY-MM-DD'
    ws['E10'] = pe; ws['E10'].number_format = 'YYYY-MM-DD'
    # 비율 정수 반올림 (6.82% → 7%) - 셀값·수식입력줄 모두 정수
    ws['F10'] = round(pm_rate * 100) / 100; ws['F10'].number_format = '0%'
    ws['D11'] = ps; ws['D11'].number_format = 'YYYY-MM-DD'
    ws['E11'] = pe; ws['E11'].number_format = 'YYYY-MM-DD'
    ws['F11'] = round(prod_rate * 100) / 100; ws['F11'].number_format = '0%' 
    ws['D17'] = STUDIO_UNIT_PRICE if include_studio else 0
    ws['E17'] = studio_hours       if include_studio else 0
    ws['E32'] = MONTHLY_SALARY
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()


# 2. 개발요청서 Excel 첨부
def gen_devreq_excel(courses, dept, month_str, year, price_tbl):
    mn = get_month_number(month_str)
    wb = openpyxl.load_workbook(os.path.join(TEMPLATE_DIR, 'tpl_devreq.xlsx'))
    ws = wb['콘텐츠개발내역']

    # 제목·작성일
    ws['A1'] = f"{mn}월 콘텐츠 개발 내역"
    ws.cell(1, 14).value = f"작성일 : {get_last_business_day(year, mn).strftime('%Y-%m-%d')}"

    # 부서행 삽입 (row2 앞에 삽입 → 기존 헤더/데이터/합계 모두 아래로)
    ws.insert_rows(2, 1)
    ws['A2'] = f"* {dept}"

    # 새 구조: 헤더=row3, 데이터=row4~, 합계=rowN
    tds = 4
    ttr = tds
    while ttr <= ws.max_row:
        if ws.cell(ttr, 1).value == '전체 합계': break
        ttr += 1
    tdc = ttr - tds

    # 행 조작 전 데이터 영역 병합 해제
    for mr in list(ws.merged_cells.ranges):
        if mr.min_row >= tds:
            ws.unmerge_cells(str(mr))

    tr_rate = price_tbl.get("1 ~ 4시간", PRICE_TRAVEL_HR)

    # 출장비 행 포함 총 행 수 계산
    rows_data = []
    for idx2, c2 in enumerate(courses):
        rows_data.append((False, c2, 0))
        ta = get_travel_for(c2, tr_rate)
        if ta > 0:
            rows_data.append((True, c2, ta))

    total_n = len(rows_data)
    if total_n > tdc:   ws.insert_rows(ttr, total_n - tdc)
    elif total_n < tdc: ws.delete_rows(tds + total_n, tdc - total_n)

    # 합계 행 설정 (루프 전에 병합 처리 → 루프 중 기존 병합 충돌 방지)
    tr = tds + total_n; er = tr - 1
    # 기존 A?:E? 병합 해제 후 새로 설정
    for mr in list(ws.merged_cells.ranges):
        if mr.min_row >= tds:
            ws.unmerge_cells(str(mr))
    ws.cell(tr, 1, "전체 합계")
    ws.merge_cells(f"A{tr}:E{tr}")
    ws.cell(tr, 6,  f"=SUM(F{tds}:F{er})")
    ws.cell(tr, 7,  f"=SUM(G{tds}:G{er})")
    ws.cell(tr, 13, f"=SUM(M{tds}:M{er})").number_format = '#,##0'
    ws.cell(tr, 14, f"=SUM(N{tds}:N{er})").number_format = '#,##0'

    course_seq = 0
    for ri, (is_travel, c, ta) in enumerate(rows_data):
        r = tds + ri
        hdr = ws.cell(tds - 1, 1)
        is_new, is_p, is_ep = classify_fmt(c.shooting_format or "")
        up = get_unit_price_for(c, price_tbl)

        if not is_travel:
            course_seq += 1
            vals = {1:course_seq, 2:"포팅" if (is_p or is_ep) else "신규",
                    3:c.course_name, 4:c.instructor, 5:"",
                    6:c.session_count or "", 7:c.chapter_count or "",
                    8:c.open_date, 9:c.shooting_format or "",
                    10:c.shooting_format or "", 11:"", 12:up,
                    13:f"=L{r}*F{r}", 14:f"=M{r}*1.1"}
        else:
            th      = c.travel_hours or 1
            td      = getattr(c, 'travel_days', None)
            loc     = getattr(c, 'location', None) or ""
            day_str = f"({td}일)" if td else ""
            loc_str = f" [{loc}]" if loc else ""
            tname   = f"{c.course_name} 출장 : {th}시간{day_str}{loc_str}"
            t_amt   = get_travel_for(c, tr_rate)  # 실제 출장비 금액
            vals    = {1:"", 2:"출장비",
                       3:tname, 4:"", 5:"",
                       6:"",   # 시간 공란 → 전체 합계 제외
                       7:"", 8:"", 9:"", 10:"", 11:"",
                       12:PRICE_TRAVEL_HR,
                       13:t_amt,              # 금액 직접값
                       14:round(t_amt*1.1)}   # 총액(VAT포함)

        for col, val in vals.items():
            try:
                cell = ws.cell(r, col, val)
                hdrc = ws.cell(tds - 1, col)
                if hdrc.border: cell.border = copy(hdrc.border)
                if col == 8 and isinstance(val, date): cell.number_format = 'YYYY-MM-DD'
                if col in (12, 13, 14): cell.number_format = '#,##0'
            except AttributeError:
                pass  # MergedCell 무시

    # 품의 시트(구 Sheet3): col A=번호, col B-H = 구분~총액 공란
    CLEAR = {2,3,4,5,6,7,8}  # B~H
    ws_pm = (wb['품의'] if '품의' in wb.sheetnames
             else wb['Sheet3'] if 'Sheet3' in wb.sheetnames else None)
    if ws_pm:
        for row in ws_pm.iter_rows():
            av = row[0].value if len(row) > 0 else None
            if av is not None and isinstance(av, (int, float)):
                for cell in row:
                    if cell.column in CLEAR:
                        try: cell.value = None
                        except: pass

    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()


# 3. 개발요청서 Word
def gen_devreq_docx(courses, dept, month_str, year, ps, pe, write_dt):
    doc   = DocxDocument(os.path.join(TEMPLATE_DIR, 'tpl_devreq.docx'))
    sj    = get_sijengil(ps)

    # 시행일, 본문 직접 치환 (각 텍스트가 별도 단락에 위치)
    _replace_doc(doc, "2026년 03월 03일", fmt_kr2(sj))   # 시행일
    _replace_doc(doc, "부서명",            dept)           # 부서명
    _replace_doc(doc, "2026년 03월 04일", fmt_kr2(ps))    # 개발기간 시작
    _replace_doc(doc, "2026년 03월 27일", fmt_kr2(pe))    # 개발기간 종료
    _replace_doc(doc, "2026년 03월 16일", fmt_kr2(write_dt))  # 서명일

    # ── row7 중첩표(번호/품목/수량/단위)에 과정 목록 채우기 ──
    cell7    = doc.tables[0].rows[7].cells[0]
    ntbl_xml = cell7._tc.findall('.//' + qn('w:tbl'))
    if ntbl_xml and courses:
        from docx.table import Table as _DTable
        from lxml import etree
        ntbl      = _DTable(ntbl_xml[0], doc)
        data_rows = ntbl.rows[1:-1]      # 헤더(0)·합계(마지막) 제외
        tpl_tr    = data_rows[0]._tr     # 행 서식 복사용

        # 빈 행 수가 부족하면 추가
        while len(data_rows) < len(courses):
            new_tr = etree.fromstring(etree.tostring(tpl_tr))
            ntbl._tbl.insert(ntbl._tbl.index(ntbl.rows[-1]._tr), new_tr)
            data_rows = ntbl.rows[1:-1]

        # 기존 데이터 행 초기화 후 입력
        total_qty = 0
        for ri, dr in enumerate(data_rows):
            for cell in dr.cells:
                for p in cell.paragraphs:
                    for r in p.runs: r.text = ""
            if ri < len(courses):
                c = courses[ri]
                is_new, is_p, is_ep = classify_fmt(c.shooting_format or "")
                unit  = "챕터" if (is_p or is_ep) else "차시"
                count = (c.chapter_count if (is_p or is_ep) else c.session_count) or 0
                total_qty += count
                for ci, val in enumerate([str(ri+1), c.course_name or "", str(count), unit]):
                    _set_cell_text(dr.cells[ci], val)

        # 합계 행 수량 채우기 (마지막 행 col 2)
        total_row = ntbl.rows[-1]
        _set_cell_text(total_row.cells[2], str(total_qty))

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# 4. 프로젝트 프로파일 Word
def gen_profile_docx(courses, dept, month_str, year, price_tbl,
                     studio_hours, include_studio, pm_rate, prod_rate,
                     customer_contact=None):
    mn = get_month_number(month_str)
    ps, pe   = calc_period(courses, year, mn)
    write_dt = get_last_business_day(year, mn)
    revenue  = calc_revenue(courses, price_tbl)
    studio_a = studio_hours * STUDIO_UNIT_PRICE if include_studio else 0
    pname    = build_project_name(courses)
    pm_pct   = round(pm_rate   * 100)
    prod_pct = round(prod_rate * 100)
    period_str = f"{fmt_short(ps)} ~ {fmt_short(pe)}"

    ns  = sum(c.session_count or 0 for c in courses if classify_fmt(c.shooting_format or "")[0])
    pc  = sum(c.chapter_count or 0 for c in courses
              if classify_fmt(c.shooting_format or "")[1] and not classify_fmt(c.shooting_format or "")[2])
    ec  = sum(c.chapter_count or 0 for c in courses if classify_fmt(c.shooting_format or "")[2])
    th  = sum((c.travel_hours or 1) for c in courses
              if c.shooting_format and "출장" in c.shooting_format)

    dparts = []
    if ns: dparts.append(f"- 신규 : {ns}차시(단가 : \\500,000, VAT별도) / 유상개발")
    if pc: dparts.append(f"- 포팅(무편집) : {pc}챕터(단가 : \\50,000, VAT별도) / 유상개발")
    if ec: dparts.append(f"- 포팅(편집) : {ec}챕터(단가 : \\160,000, VAT별도) / 유상개발")
    if th: dparts.append(f"- 출장비 : {th}시간(단가 : \\100,000, VAT별도)")
    new_detail = ("1. 사업내용\n(1) 한공회 콘텐츠 개발\n"
                  + "\n".join(dparts)
                  + f"\n(2) 개발기간 : {fmt_kr(ps)} ~ {fmt_kr(pe)}")

    cn = customer_contact.contact_name if customer_contact else ""
    cp = customer_contact.phone        if customer_contact else ""
    ce = customer_contact.email        if customer_contact else ""

    doc = DocxDocument(os.path.join(TEMPLATE_DIR, 'tpl_profile.docx'))
    t0,t1,t2,t3,t4 = doc.tables[0],doc.tables[1],doc.tables[2],doc.tables[3],doc.tables[4]

    # ── 새 템플릿 플레이스홀더 치환 ──
    # 표0: 작성일
    _replace_doc(doc, "작성일 표시",   fmt_kr(write_dt))
    # 표1: 프로젝트명·기간·계약금액
    _replace_doc(doc, "한공회 콘텐츠 개발(신규00차시, 포팅00챕터)", pname)
    _replace_doc(doc, "프로젝트 기간",          f"{fmt_kr(ps)} ~ {fmt_kr(pe)}")
    _replace_doc(doc, "계약금액 : 금액표시",    f"계약금액 : \u20a9{revenue:,}")
    _replace_doc(doc, "계약금액 표시",          f"\u20a9{revenue:,}")
    # 표1: 세부내역·담당자
    _clear_cell_and_set(t1.rows[8].cells[1], new_detail)
    _set_cell_text(t1.rows[12].cells[1], cn)
    _set_cell_text(t1.rows[12].cells[2], cp)
    _set_cell_text(t1.rows[12].cells[5], ce)
    # 표2: 참여인력 (참여기간·참여율 직접 설정)
    _set_cell_text(t2.rows[2].cells[3], period_str)  # PM 참여기간
    _set_cell_text(t2.rows[3].cells[3], period_str)  # PROD 참여기간
    _set_cell_text(t2.rows[2].cells[4], f"{pm_pct}%")   # PM 참여율
    _set_cell_text(t2.rows[3].cells[4], f"{prod_pct}%") # PROD 참여율

    if include_studio and studio_hours > 0:
        _set_cell_text(t4.rows[2].cells[1], str(studio_hours))
        _set_cell_text(t4.rows[2].cells[2], f"{STUDIO_UNIT_PRICE:,}")
        _set_cell_text(t4.rows[2].cells[3], f"{studio_hours*STUDIO_UNIT_PRICE:,}")
    else:
        for ci in range(5): _set_cell_text(t4.rows[2].cells[ci], "")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# 전체 ZIP 생성
def generate_all(courses, dept, month_str, year, price_tbl,
                 studio_hours, include_studio, customer_contact=None):
    mn       = get_month_number(month_str)
    ps, pe   = calc_period(courses, year, mn)
    write_dt = get_last_business_day(year, mn)
    revenue  = calc_revenue(courses, price_tbl)
    studio_a = studio_hours * STUDIO_UNIT_PRICE if include_studio else 0
    # 출장비 포함 전체 매출 기준으로 투입비율 계산 (출장비 1차시당 1시간 반영)
    pm_rate, prod_rate = adjust_rates(revenue, studio_a, ps, pe, courses=courses)

    mm_str = f"{mn:02d}월"
    mmdd   = write_dt.strftime('%m%d')   # 작성일 MMDD (예: 0430)

    pnl      = gen_pnl_excel(courses, dept, month_str, year, price_tbl,
                              studio_hours, include_studio, pm_rate, prod_rate)
    rx       = gen_devreq_excel(courses, dept, month_str, year, price_tbl)
    rd       = gen_devreq_docx(courses, dept, month_str, year, ps, pe, write_dt)
    prof     = gen_profile_docx(courses, dept, month_str, year, price_tbl,
                                 studio_hours, include_studio,
                                 pm_rate, prod_rate, customer_contact)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"손익분석서_한국공인회계사회_{mm_str}_V1.0_{mmdd}_{dept}.xlsx",                  pnl)
        zf.writestr(f"프로젝트프로파일_한국공인회계사회_{mm_str}_V0.1_{mmdd}_{dept}.docx",             prof)
        zf.writestr(f"한공회_{year}년{mn:02d}월분_컨텐츠개발요청서_{dept}.docx",                       rd)
        zf.writestr(f"한공회_{year}년{mn:02d}월분_컨텐츠개발요청서_제출시첨부사항_{dept}.xlsx",         rx)
    return buf.getvalue()
